from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

RED  = RGBColor(0xDC, 0x26, 0x26)
DARK = RGBColor(0x0F, 0x17, 0x2A)

def header(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = RED
    r.font.name = "Calibri"
    rPr = r._r.get_or_add_rPr()
    sp = OxmlElement("w:spacing"); sp.set(qn("w:val"), "60"); rPr.append(sp)

def line(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(11)
    r.font.color.rgb = DARK
    r.font.name = "Calibri"

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "2")
    bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), "CBD5E1")
    pBdr.append(bot); pPr.append(pBdr)

# ── Заголовок ──────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Речь · Модуль аналитики и адаптивности · Gradus")
r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RED; r.font.name = "Calibri"

divider(doc)

# ── ВСТУПЛЕНИЕ ────────────────────────────────────────────────────────────
header(doc, "ВСТУПЛЕНИЕ")
line(doc, "Добрый день, уважаемая комиссия!")
line(doc, "В рамках группового проекта моим модулем является проектирование и разработка модуля аналитики и адаптивности — ключевого компонента, который позволяет реализовать поставленные задачи работы.")

divider(doc)

# ── ПРОБЛЕМА ──────────────────────────────────────────────────────────────
header(doc, "ПРОБЛЕМА")
line(doc, "Большинство существующих платформ — это просто «хранилище» материалов с проверкой тестов и шаблонных ответов. Платформы не понимают, как учится студент.")
line(doc, "Мой модуль решает именно эту проблему: он собирает глубокую аналитику действий и на её основе даёт полную обратную связь и может изменять обучение под каждого пользователя.")

divider(doc)

# ── ГЛАВНЫЕ ЗАДАЧИ ────────────────────────────────────────────────────────
header(doc, "ГЛАВНЫЕ ЗАДАЧИ МОДУЛЯ")
line(doc, "Первое — сбор детальных метрик: время выполнения, количество попыток, сложность заданий.")
line(doc, "Второе — расчёт интегрального прогресса и прогноза завершения курса.")
line(doc, "Третье — адаптивная проверка заданий с частичным оцениванием.")
line(doc, "И четвёртое — формирование персонализированных AI-инсайтов и рекомендаций.")

divider(doc)

# ── ТЕХНОЛОГИЧЕСКАЯ ОСНОВА ────────────────────────────────────────────────
header(doc, "ТЕХНОЛОГИЧЕСКАЯ ОСНОВА")
line(doc, "Бэкенд модуля реализован на Python с использованием фреймворка FastAPI. Почему FastAPI? Он обеспечивает высокую производительность за счёт асинхронности, что критично при одновременной работе многих пользователей. Встроенная валидация через Pydantic гарантирует чистоту входящих данных.")
line(doc, "Для хранения данных используется PostgreSQL — надёжная реляционная база с поддержкой JSON для гибких сценариев."  )

divider(doc)

# ── ЗАКЛЮЧЕНИЕ ────────────────────────────────────────────────────────────
header(doc, "ЗАКЛЮЧЕНИЕ")
line(doc, "Разработанный модуль аналитики и адаптивности полностью работоспособен. Он собирает метрики, динамически подстраивает рекомендации и выдаёт AI-инсайты. Тестирование подтвердило корректность всех сценариев.")
line(doc, "Таким образом, задача создания платформы, которая не просто хранит контент, а понимает и адаптируется под студента в реальном времени, успешно решена.", bold=True)

# ── Сохранить ─────────────────────────────────────────────────────────────
out = r"C:\Users\yaros\Desktop\stepashka\шпаргалка_Виктория.docx"
doc.save(out)
print(f"Saved: {out}")
